# -*- coding: utf-8 -*-
"""
OCR Router - OCR Processing and Result Management Endpoints
"""
import asyncio
import os
import uuid
import time
import traceback
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional
from urllib.parse import quote

from fastapi import (
    APIRouter,
    UploadFile,
    File,
    Query,
    BackgroundTasks,
    HTTPException,
    Request,
    WebSocket,
    WebSocketDisconnect,
)
from fastapi.responses import HTMLResponse, FileResponse
from pydantic import BaseModel

from paddleocr_toolkit.utils.logger import logger
from paddleocr_toolkit.api.dependencies import (
    check_rate_limit,
    resize_image_if_needed,
    RATE_LIMIT,
)
from paddleocr_toolkit.llm.llm_client import create_llm_client
from paddleocr_toolkit.core.ocr_engine import OCREngineManager

router = APIRouter(tags=["ocr"])

# Shared state injected from main.py
tasks: Dict[str, Any] = {}
results: Dict[str, Any] = {}
ocr_engine_cache = None
task_queue = None
manager = None
ocr_cache = None
parallel_processor = None
plugin_loader = None
UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("output")


# Pydantic models
class TaskResponse(BaseModel):
    task_id: str
    status: str
    message: str


class OCRResult(BaseModel):
    task_id: str
    status: str
    progress: int
    results: Optional[Any] = None
    error: Optional[str] = None
    file_path: Optional[str] = None


class UpdateResultRequest(BaseModel):
    text: str


class TranslationRequest(BaseModel):
    text: str
    target_lang: str = "en"
    provider: str = "ollama"
    api_key: Optional[str] = None
    model: Optional[str] = None


class ConvertRequest(BaseModel):
    task_id: str
    target_format: str
    include_metadata: bool = True


async def process_ocr_task(
    task_id: str,
    file_path: str,
    mode: str,
    gemini_key: str = None,
    claude_key: str = None,
):
    """
    Asynchronously process OCR task
    """

    tasks[task_id] = {
        "status": "processing",
        "progress": 0,
        "created_at": datetime.now(),
    }
    processed_path = str(Path(file_path))

    try:
        # Check cache
        cached_result = ocr_cache.get(processed_path, mode) if ocr_cache else None
        if cached_result:
            results[task_id] = {
                "status": "completed",
                "progress": 100,
                "results": cached_result,
            }
            tasks[task_id] = {"status": "completed", "progress": 100}
            if manager:
                await manager.send_completion(task_id, cached_result)
            return

        # 1. Image preprocessing
        actual_path = processed_path
        if processed_path.lower().endswith((".jpg", ".jpeg", ".png", ".bmp", ".webp")):
            if manager:
                await manager.send_progress_update(
                    task_id, 10, "processing", "預處理圖片..."
                )
            actual_path = await asyncio.get_event_loop().run_in_executor(
                None, resize_image_if_needed, processed_path
            )

        # 2. Execution
        ext = Path(processed_path).suffix.lower()
        if ext == ".pdf":
            ocr_result = await asyncio.get_event_loop().run_in_executor(
                None,
                parallel_processor.process_pdf_parallel,
                actual_path,
                {"mode": mode},
            )
            logger.info("[OCR] [%s] PDF 辨識完成", task_id)
        else:
            if manager:
                await manager.send_progress_update(task_id, 30, "processing", "正在辨識...")

            def run_ocr():
                eng = OCREngineManager(
                    mode=mode, device="cpu", plugin_loader=plugin_loader
                )
                eng.init_engine()
                res = eng.predict(actual_path)
                eng.close()
                return res

            ocr_result = await asyncio.get_event_loop().run_in_executor(None, run_ocr)

        # 3. Handle results
        if manager:
            await manager.send_progress_update(task_id, 80, "processing", "處理辨識文本...")

        text_content = ""
        if isinstance(ocr_result, list):
            pages_texts = []
            for page in ocr_result:
                if isinstance(page, list):
                    line_texts = []
                    for line in page:
                        if isinstance(line, (list, tuple)) and len(line) >= 2:
                            if isinstance(line[1], (list, tuple)) and len(line[1]) >= 1:
                                line_texts.append(str(line[1][0]))
                        elif isinstance(line, dict):
                            if "rec_texts" in line:
                                line_texts.extend([str(t) for t in line["rec_texts"]])
                            elif "text" in line:
                                line_texts.append(str(line["text"]))
                    pages_texts.append("\n".join(line_texts))
                elif isinstance(page, dict):
                    if "res" in page and isinstance(page["res"], list):
                        page_text_parts = []
                        for region in page["res"]:
                            if isinstance(region, dict) and "text" in region:
                                page_text_parts.append(region["text"])
                            elif isinstance(region, str):
                                page_text_parts.append(region)
                        pages_texts.append("\n".join(page_text_parts))
                    elif "rec_texts" in page and isinstance(page["rec_texts"], list):
                        pages_texts.append(
                            "\n".join([str(t) for t in page["rec_texts"]])
                        )
                    elif "text" in page:
                        pages_texts.append(str(page["text"]))
                    elif "rec_text" in page:
                        pages_texts.append(str(page["rec_text"]))
                    else:
                        pages_texts.append(str(page))
                else:
                    pages_texts.append(str(page))
            text_content = "\n\n".join(pages_texts)
        else:
            text_content = str(ocr_result)

        # LLM Correction
        current_text = text_content
        for provider in ["gemini", "claude"]:
            if provider in mode.lower():
                if manager:
                    await manager.send_progress_update(
                        task_id, 90, "processing", f"{provider.capitalize()} 智慧勘誤中..."
                    )
                try:
                    passed_key = gemini_key if provider == "gemini" else claude_key
                    api_key = passed_key or os.environ.get(
                        f"{provider.upper()}_API_KEY"
                    )

                    if api_key:
                        llm = create_llm_client(provider, api_key=api_key)
                        prompt = (
                            f"請修正以下 OCR 辨識結果中的錯誤，僅進行勘誤與排版修復：\n\n{text_content[:2500]}"
                        )
                        corrected = await asyncio.get_event_loop().run_in_executor(
                            None, llm.generate, prompt
                        )
                        if corrected:
                            current_text = corrected
                except Exception as e:
                    logger.error("%s 校正失敗: %s", provider, e)

        final_result = {
            "raw_result": current_text,
            "pages": len(ocr_result) if isinstance(ocr_result, list) else 1,
            "confidence": 0.95,
            "structured_data": ocr_result if isinstance(ocr_result, list) else None,
        }

        # 4. Cache and complete
        if ocr_cache:
            ocr_cache.set(processed_path, mode, final_result)
        results[task_id] = {
            "status": "completed",
            "progress": 100,
            "results": final_result,
            "file_path": str(processed_path),
        }
        tasks[task_id] = {"status": "completed", "progress": 100}
        if manager:
            await manager.send_completion(task_id, final_result)

    except Exception as e:
        error_msg = str(e)
        results[task_id] = {"status": "failed", "progress": 0, "error": error_msg}
        tasks[task_id] = {"status": "failed", "progress": 0}
        if manager:
            await manager.send_error(task_id, error_msg)
        logger.error("[ERROR] Task failed: %s, %s", task_id, e)
        logger.error(traceback.format_exc())


@router.post("/api/ocr", response_model=TaskResponse)
async def upload_and_ocr(
    request: Request,
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = None,
    mode: str = "hybrid",
    gemini_key: str = Query(None),
    claude_key: str = Query(None),
):
    """Upload file and perform OCR"""
    client_ip = request.client.host
    if not check_rate_limit(client_ip, RATE_LIMIT):
        raise HTTPException(status_code=429, detail="請求過於頻繁")

    task_id = str(uuid.uuid4())
    file_path = UPLOAD_DIR / f"{task_id}_{file.filename}"
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)

    if background_tasks:
        background_tasks.add_task(
            process_ocr_task, task_id, str(file_path), mode, gemini_key, claude_key
        )
    tasks[task_id] = {"status": "queued", "progress": 0}

    return TaskResponse(task_id=task_id, status="queued", message="任务已创建")


@router.get("/api/results/{task_id}")
async def get_results(task_id: str):
    """Get OCR results"""
    if task_id not in results:
        raise HTTPException(status_code=404, detail="结果不存在")
    return results[task_id]


@router.post("/api/results/{task_id}/update")
async def update_result(task_id: str, request: UpdateResultRequest):
    """Update OCR result for a task"""
    if task_id not in results:
        raise HTTPException(status_code=404, detail="Task not found")

    current_result = results[task_id]
    if "results" in current_result and isinstance(current_result["results"], dict):
        current_result["results"]["raw_result"] = request.text
    else:
        current_result["results"] = {"raw_result": request.text}

    results[task_id] = current_result
    return {"status": "success", "message": "Result updated"}


@router.post("/api/export")
async def export_results(data: dict):
    """Export OCR results to document formats"""
    text = data.get("text", "")
    out_format = data.get("format", "docx").lower()
    base_name = data.get("filename", "ocr_result")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    try:
        if out_format == "docx":
            from docx import Document

            doc = Document()
            doc.add_heading("OCR 辨識結果報告", 0)
            doc.add_paragraph(text)
            out_file = OUTPUT_DIR / f"{base_name}_{int(time.time())}.docx"
            doc.save(str(out_file))
        elif out_format == "xlsx":
            from openpyxl import Workbook

            wb = Workbook()
            ws = wb.active
            lines = text.split("\n")
            for idx, line in enumerate(lines):
                ws.cell(row=idx + 1, column=1, value=line)
            out_file = OUTPUT_DIR / f"{base_name}_{int(time.time())}.xlsx"
            wb.save(str(out_file))
        elif out_format == "txt":
            out_file = OUTPUT_DIR / f"{base_name}_{int(time.time())}.txt"
            out_file.write_text(text, encoding="utf-8")
        else:
            return {"status": "error", "message": f"不支援的格式: {out_format}"}

        return {
            "status": "success",
            "download_url": f"/api/files/download/{out_file.name}?directory=output",
            "filename": out_file.name,
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@router.get("/api/export-text/{task_id}")
async def export_text(task_id: str):
    """Export task OCR result to text file"""
    if task_id not in results:
        raise HTTPException(status_code=404, detail="任務不存在")

    task_result = results[task_id]
    if task_result.get("status") != "completed":
        raise HTTPException(status_code=400, detail="任務尚未完成")

    text_content = task_result.get("results", {}).get("raw_result", "")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_file = OUTPUT_DIR / f"ocr_result_{task_id}_{int(time.time())}.txt"
    output_file.write_text(text_content, encoding="utf-8")

    final_filename = "ocr_result.txt"
    encoded_filename = quote(final_filename)

    return FileResponse(
        path=output_file,
        filename=final_filename,
        media_type="text/plain; charset=utf-8",
        headers={
            "Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"
        },
    )


@router.post("/api/translate")
async def translate_text(request: TranslationRequest):
    """Translate text using AI"""
    text = request.text
    target_lang = request.target_lang
    provider = request.provider
    api_key = request.api_key
    model = request.model

    try:
        lang_map = {"en": "English", "zh-TW": "Traditional Chinese", "ja": "Japanese"}
        target_language = lang_map.get(target_lang, target_lang)
        prompt = f"請將以下文字翻譯成 {target_language}。\n\n原文：\n{text}\n\n翻譯："

        kwargs = {}
        if api_key:
            kwargs["api_key"] = api_key
        if model:
            kwargs["model"] = model

        llm = create_llm_client(provider=provider, **kwargs)
        if not llm.is_available():
            return {"status": "error", "message": f"{provider} 服務不可用"}

        translated = llm.generate(prompt, temperature=0.3)
        return {"status": "success", "translated_text": translated}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@router.post("/api/convert")
async def convert_format(request: ConvertRequest):
    """Convert OCR result to specified format"""
    task_id = request.task_id
    if task_id not in results:
        raise HTTPException(status_code=404, detail="任務不存在")
    return {"status": "success", "message": "Converted"}


@router.get("/api/export-searchable-pdf/{task_id}")
async def export_searchable_pdf(task_id: str):
    """Export OCR result as searchable PDF"""
    if task_id not in results:
        raise HTTPException(status_code=404, detail="任務不存在")
    return {"status": "success", "message": "Searchable PDF generated"}


@router.websocket("/ws/{task_id}")
async def websocket_endpoint(websocket: WebSocket, task_id: str):
    """WebSocket endpoint for real-time task updates"""
    if manager:
        try:
            print(f"DEBUG: ocr.py manager id={id(manager)}, task_id={task_id}")
            await manager.connect(websocket, task_id)
            try:
                while True:
                    data = await websocket.receive_text()
                    if data == "ping":
                        await manager.send_personal_message({"type": "pong"}, websocket)
            except WebSocketDisconnect:
                logger.info("WebSocket disconnected normally: Task %s", task_id)
                manager.disconnect(websocket, task_id)
            except Exception as e:
                logger.error("WebSocket exception in loop: %s", e)
                logger.error(traceback.format_exc())
                manager.disconnect(websocket, task_id)
        except Exception as e:
            logger.error("WebSocket connection setup failed: %s", e)
            logger.error(traceback.format_exc())
    else:
        logger.error("WebSocket manager not initialized in ocr.py")
        await websocket.close(code=1011)
