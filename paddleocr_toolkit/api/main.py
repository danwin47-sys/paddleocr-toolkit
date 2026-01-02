#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FastAPI后端 - Web界面
v1.2.0新增 - REST API服务
"""
from __future__ import annotations  # Python 3.8 兼容性

import asyncio
import os
import time
import uuid
from collections import defaultdict
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Optional
from urllib.parse import quote

from fastapi import (
    BackgroundTasks,
    FastAPI,
    File,
    HTTPException,
    Query,
    Request,
    UploadFile,
    WebSocket,
    WebSocketDisconnect,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel



from paddleocr_toolkit.api.file_manager import router as file_router
from paddleocr_toolkit.api.websocket_manager import manager
from paddleocr_toolkit.core.ocr_engine import OCREngineManager
from paddleocr_toolkit.plugins.loader import PluginLoader

# 找到 web 目錄
WEB_DIR = Path(__file__).parent.parent.parent / "web"
STATIC_DIR = WEB_DIR  # 目前 index.html 直接在 web 根目錄

app = FastAPI(
    title="PaddleOCR Toolkit API",
    description="專業級 OCR 文件處理 API",
    version="3.3.0",
    docs_url="/docs",
    redoc_url="/redoc",
)


# Global Exception Handler
@app.exception_handler(Exception)
async def global_exception_handler(request, exc):
    import traceback
    error_detail = traceback.format_exc()
    print(f"[CRITICAL ERROR] {exc}")
    print(error_detail)
    return JSONResponse(
        status_code=500,
        content={"status": "error", "message": f"伺服器內部錯誤: {str(exc)}", "detail": error_detail}
    )


# Request Logging Middleware
@app.middleware("http")
async def log_requests(request, call_next):
    import time
    start_time = time.time()
    path = request.url.path
    method = request.method
    print(f"[HTTP] {method} {path} - 收到請求")
    
    try:
        response = await call_next(request)
        duration = time.time() - start_time
        print(f"[HTTP] {method} {path} - 完成處理 ({duration:.2f}s) - Status: {response.status_code}")
        return response
    except Exception as e:
        print(f"[HTTP] {method} {path} - 處理時發生異常: {e}")
        import traceback
        print(traceback.format_exc())
        raise e

# 掛載路由
app.include_router(file_router)

# 靜態檔案服務 (如果目錄存在)
if WEB_DIR.exists():
    app.mount("/web", StaticFiles(directory=str(WEB_DIR)), name="web")

# CORS設定
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],  # 允許前端讀取自訂 header
)

# 任務儲存
tasks = {}
results = {}
batches = {}  # 批量處理追蹤

# 簡單速率限制器
rate_limits = defaultdict(list)
RATE_LIMIT = 10  # 每分鐘請求數
RATE_LIMIT_BATCH = 3  # 批量處理每分鐘請求數
RATE_WINDOW = 60  # 時間窗口（秒）

# 配置
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(exist_ok=True)
PLUGIN_DIR = Path("paddleocr_toolkit/plugins")

# 初始化插件載入器
plugin_loader = PluginLoader(str(PLUGIN_DIR))
plugin_loader.load_all_plugins()

# 圖片大小限制 (避免 OCR 記憶體不足)
MAX_IMAGE_SIDE = 2500  # 像素

# 全域 OCR 引擎快取
ocr_engine_cache = None


def check_rate_limit(client_ip: str, limit: int = RATE_LIMIT) -> bool:
    """
    檢查速率限制
    
    Args:
        client_ip: 客戶端 IP
        limit: 請求限制（預設 10/分鐘）
    
    Returns:
        True 如果未超限，False 如果超限
    """
    now = time.time()
    
    # 清理舊記錄
    rate_limits[client_ip] = [
        t for t in rate_limits[client_ip] 
        if now - t < RATE_WINDOW
    ]
    
    # 檢查是否超限
    if len(rate_limits[client_ip]) >= limit:
        return False
    
    # 記錄請求
    rate_limits[client_ip].append(now)
    return True


async def cleanup_old_tasks():
    """
    定期清理舊任務，防止記憶體洩漏和磁碟空間耗盡
    每小時清理 24 小時前的任務和檔案
    """
    while True:
        try:
            await asyncio.sleep(3600)  # 每小時執行一次
            
            cutoff = datetime.now() - timedelta(hours=24)
            removed_tasks = 0
            removed_results = 0
            removed_files = 0
            
            # 清理舊任務
            for task_id in list(tasks.keys()):
                task = tasks.get(task_id, {})
                task_time = task.get('created_at')
                
                if task_time and isinstance(task_time, datetime) and task_time < cutoff:
                    del tasks[task_id]
                    removed_tasks += 1
            
            # 清理舊結果
            for task_id in list(results.keys()):
                if task_id not in tasks:
                    del results[task_id]
                    removed_results += 1
            
            # 清理舊檔案（上傳和輸出）
            cutoff_timestamp = cutoff.timestamp()
            
            # 清理上傳目錄
            for file_path in UPLOAD_DIR.glob("*"):
                if file_path.is_file():
                    try:
                        if file_path.stat().st_mtime < cutoff_timestamp:
                            file_path.unlink()
                            removed_files += 1
                    except Exception as e:
                        print(f"⚠️ 無法刪除檔案 {file_path}: {e}")
           
            # 清理輸出目錄
            for file_path in OUTPUT_DIR.glob("*"):
                if file_path.is_file():
                    try:
                        if file_path.stat().st_mtime < cutoff_timestamp:
                            file_path.unlink()
                            removed_files += 1
                    except Exception as e:
                        print(f"⚠️ 無法刪除檔案 {file_path}: {e}")
            
            if removed_tasks > 0 or removed_results > 0 or removed_files > 0:
                print(f"🧹 清理完成: 移除 {removed_tasks} 個舊任務, {removed_results} 個舊結果, {removed_files} 個舊檔案")
        except Exception as e:
            print(f"⚠️ 清理任務時發生錯誤: {e}")


@app.on_event("startup")
async def startup_event():
    """
    應用啟動時預載 OCR 引擎並啟動清理任務
    """
    global ocr_engine_cache
    print("=" * 60)
    print("🚀 正在預載 OCR 引擎...")
    print("=" * 60)
    try:
        # 預載基礎模式引擎（最常用）
        ocr_engine_cache = OCREngineManager(mode="basic", device="cpu", plugin_loader=plugin_loader)
        ocr_engine_cache.init_engine()
        print("✅ OCR 引擎預載完成 (Basic 模式)")
        print("   首次 OCR 請求將直接使用預載引擎，無需等待模型載入")
    except Exception as e:
        print(f"⚠️ OCR 引擎預載失敗: {e}")
        ocr_engine_cache = None
    
    # 啟動定期清理任務
    asyncio.create_task(cleanup_old_tasks())
    print("🧹 已啟動定期任務清理（每小時執行）")
    print("=" * 60)


def resize_image_if_needed(file_path: str, max_side: int = MAX_IMAGE_SIDE) -> str:
    """
    檢測並縮小大圖片以避免 OCR 記憶體問題

    Args:
        file_path: 圖片路徑
        max_side: 最大邊長（預設 2500px）

    Returns:
        處理後的圖片路徑（縮小後的新檔案或原檔案）
    """
    try:
        from PIL import Image

        with Image.open(file_path) as img:
            width, height = img.size
            max_dim = max(width, height)

            if max_dim <= max_side:
                print(f"圖片大小 {width}x{height} 在限制內，不需要縮小")
                return file_path

            # 計算縮放比例
            scale = max_side / max_dim
            new_width = int(width * scale)
            new_height = int(height * scale)

            print(f"圖片太大 ({width}x{height})，縮小為 {new_width}x{new_height}")

            # 縮小圖片
            resized_img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)

            # 儲存到新檔案
            path = Path(file_path)
            new_path = path.parent / f"{path.stem}_resized{path.suffix}"
            resized_img.save(str(new_path), quality=95)

            print(f"縮小後圖片已儲存: {new_path}")
            return str(new_path)

    except Exception as e:
        print(f"縮小圖片時發生錯誤: {e}，使用原始圖片")
        return file_path


class OCRRequest(BaseModel):
    """OCR请求模型"""

    mode: str = "hybrid"
    dpi: int = 200
    lang: str = "ch"


class TaskResponse(BaseModel):
    """任务响应模型"""

    task_id: str
    status: str
    message: str


class OCRResult(BaseModel):
    """OCR结果模型"""

    task_id: str
    status: str
    progress: float
    results: Optional[Any] = None
    error: Optional[str] = None


from paddleocr_toolkit.core.ocr_cache import ocr_cache
from paddleocr_toolkit.processors.parallel_pdf_processor import ParallelPDFProcessor

# 初始化平行處理器
parallel_processor = ParallelPDFProcessor()


async def process_ocr_task(
    task_id: str,
    file_path: str,
    mode: str,
    gemini_key: str = None,
    claude_key: str = None,
):
    """
    非同步處理 OCR 任務
    """
    # 詳細日誌：任務開始
    print("=" * 60)
    print(f"[OCR] 開始處理任務")
    print(f"[OCR] 任務 ID: {task_id}")
    print(f"[OCR] 檔案路徑: {file_path}")
    print(f"[OCR] OCR 模式: {mode}")
    print(f"[OCR] Gemini Key: {'已提供' if gemini_key else '未提供'}")
    print(f"[OCR] Claude Key: {'已提供' if claude_key else '未提供'}")
    print("=" * 60)
    
    tasks[task_id] = {
        "status": "processing",
        "progress": 0,
        "created_at": datetime.now()
    }
    processed_path = str(Path(file_path))
    print(f"[OCR] 處理路徑: {processed_path}")

    try:
        # ... (快取與處理邏輯)
        # 0. 檢查快取
        await manager.send_progress_update(task_id, 5, "processing", "檢查快取...")
        cached_result = ocr_cache.get(processed_path, mode)
        if cached_result:
            await manager.send_progress_update(task_id, 100, "completed", "快取命中")
            results[task_id] = {
                "status": "completed",
                "progress": 100,
                "results": cached_result,
            }
            tasks[task_id] = {"status": "completed", "progress": 100}
            await manager.send_completion(task_id, cached_result)
            return

        # 1. 預處理圖片
        actual_path = processed_path
        if processed_path.lower().endswith((".jpg", ".jpeg", ".png", ".bmp", ".webp")):
            await manager.send_progress_update(task_id, 10, "processing", "預處理圖片...")
            actual_path = await asyncio.to_thread(
                resize_image_if_needed, processed_path
            )

        # 2. 執行預測
        ext = Path(processed_path).suffix.lower()
        if ext == ".pdf":
            # PDF 使用並行處理器
            await manager.send_progress_update(
                task_id, 30, "processing", "PDF 並行辨識中..."
            )
            ocr_result = await asyncio.to_thread(
                parallel_processor.process_pdf_parallel, actual_path, {"mode": mode}
            )
        else:
            # 標準單點辨識
            await manager.send_progress_update(task_id, 30, "processing", "正在辨識...")

            def run_ocr():
                eng = OCREngineManager(
                    mode=mode, device="cpu", plugin_loader=plugin_loader
                )
                eng.init_engine()
                res = eng.predict(actual_path)
                eng.close()
                return res

            ocr_result = await asyncio.to_thread(run_ocr)

        # 3. 處理與校正結果
        await manager.send_progress_update(task_id, 80, "processing", "處理辨識文本...")

        # 提取文字內容
        text_content = ""
        if isinstance(ocr_result, list):
            pages_texts = []
            for page in ocr_result:
                if isinstance(page, list):
                    # PaddleOCR 標準格式: [[[box], [text, score]], ...]
                    line_texts = []
                    for line in page:
                        if isinstance(line, (list, tuple)) and len(line) >= 2:
                            if isinstance(line[1], (list, tuple)) and len(line[1]) >= 1:
                                line_texts.append(str(line[1][0]))
                        elif isinstance(line, dict):
                            # PaddleX 字典格式 (在列表內)
                            if "rec_texts" in line:
                                line_texts.extend([str(t) for t in line["rec_texts"]])
                            elif "text" in line:
                                line_texts.append(str(line["text"]))

                    pages_texts.append("\n".join(line_texts))

                elif isinstance(page, dict):
                    # PPStructure or newer PaddleOCR format
                    # 1. Check for 'res' (PPStructure standard)
                    if "res" in page and isinstance(page["res"], list):
                        page_text_parts = []
                        for region in page["res"]:
                            if isinstance(region, dict) and "text" in region:
                                page_text_parts.append(region["text"])
                            elif isinstance(region, str):
                                page_text_parts.append(region)
                        pages_texts.append("\n".join(page_text_parts))

                    # 2. Check for 'rec_texts' (PaddleOCR Direct Result / Parallel Processor)
                    elif "rec_texts" in page and isinstance(page["rec_texts"], list):
                        pages_texts.append(
                            "\n".join([str(t) for t in page["rec_texts"]])
                        )

                    # 3. Check for 'text' or 'rec_text' (Single String)
                    elif "text" in page:
                        pages_texts.append(str(page["text"]))
                    elif "rec_text" in page:
                        pages_texts.append(str(page["rec_text"]))

                    # Fallback
                    else:
                        pages_texts.append(str(page))
                else:
                    pages_texts.append(str(page))
            text_content = "\n\n".join(pages_texts)
        else:
            text_content = str(ocr_result)

        # 語義校正 (Gemini / Claude)
        current_text = text_content
        for provider in ["gemini", "claude"]:
            if provider in mode.lower():
                await manager.send_progress_update(
                    task_id, 90, "processing", f"{provider.capitalize()} 智慧勘誤中..."
                )
                try:
                    import os

                    from paddleocr_toolkit.llm.llm_client import create_llm_client

                    # 優先使用傳入的金鑰，次之使用環境變數
                    passed_key = gemini_key if provider == "gemini" else claude_key
                    api_key = passed_key or os.environ.get(
                        f"{provider.upper()}_API_KEY"
                    )

                    if api_key:
                        llm = create_llm_client(provider, api_key=api_key)
                        prompt = (
                            f"請修正以下 OCR 辨識結果中的錯誤，僅進行勘誤與排版修復：\n\n{text_content[:2500]}"
                        )
                        corrected = await asyncio.to_thread(llm.generate, prompt)
                        if corrected:
                            current_text = corrected
                except Exception as e:
                    print(f"{provider} 校正失敗: {e}")

        final_result = {
            "raw_result": current_text,
            "pages": len(ocr_result) if isinstance(ocr_result, list) else 1,
            "confidence": 0.95,
        }

        # 4. 存入快取與完成
        ocr_cache.set(processed_path, mode, final_result)
        results[task_id] = {
            "status": "completed",
            "progress": 100,
            "results": final_result,
        }
        tasks[task_id] = {"status": "completed", "progress": 100}
        await manager.send_completion(task_id, final_result)

    except Exception as e:
        error_msg = str(e)
        results[task_id] = {"status": "failed", "progress": 0, "error": error_msg}
        tasks[task_id] = {"status": "failed", "progress": 0}
        results[task_id] = {
            "status": "failed",
            "progress": 0,
            "error": error_msg
        }
        await manager.send_error(task_id, error_msg)
        
        # 詳細錯誤日誌
        print("=" * 60)
        print(f"[ERROR] 任務失敗: {task_id}")
        print(f"[ERROR] 錯誤訊息: {e}")
        print(f"[ERROR] 完整 Traceback:")
        import traceback
        traceback.print_exc()
        print("=" * 60)


@app.get("/", response_class=HTMLResponse)
async def root():
    """提供用戶友好的 Web 介面"""
    index_file = WEB_DIR / "index.html"
    if index_file.exists():
        return HTMLResponse(content=index_file.read_text(encoding="utf-8"))
    return HTMLResponse(
        content="""
        <h1>PaddleOCR Toolkit API</h1>
        <p>Version: 1.2.0</p>
        <p><a href="/docs">API 文件</a></p>
    """
    )


@app.post("/api/ocr", response_model=TaskResponse)
async def upload_and_ocr(
    request: Request,
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = None,  # type: ignore[assignment]
    mode: str = "hybrid",
    gemini_key: str = Query(None),
    claude_key: str = Query(None),
):
    """
    上传文件并进行OCR处理

    Args:
        file: 上传的文件
        background_tasks: 后台任务
        mode: OCR模式

    Returns:
        任务ID和状态
    """
    # 速率限制檢查
    client_ip = request.client.host
    if not check_rate_limit(client_ip, RATE_LIMIT):
        raise HTTPException(
            status_code=429,
            detail="請求過於頻繁，請稍後再試（限制：10 次/分鐘）"
        )
    
    # 生成任务ID
    task_id = str(uuid.uuid4())
    print(f"[UPLOAD] 收到上傳請求")
    print(f"[UPLOAD] 任務 ID: {task_id}")
    print(f"[UPLOAD] 檔案名稱: {file.filename}")
    print(f"[UPLOAD] OCR 模式: {mode}")

    # 保存文件
    file_path = UPLOAD_DIR / f"{task_id}_{file.filename}"
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    print(f"[UPLOAD] 檔案已儲存: {file_path}")
    print(f"[UPLOAD] 檔案大小: {len(content)} bytes")

    # 创建后台任务
    background_tasks.add_task(
        process_ocr_task, task_id, str(file_path), mode, gemini_key, claude_key
    )
    print(f"[UPLOAD] 背景任務已排程: {task_id}")

    # 初始化任务状态
    tasks[task_id] = {"status": "queued", "progress": 0}
    print(f"[UPLOAD] 任務狀態已初始化: queued")

    return TaskResponse(task_id=task_id, status="queued", message="任务已创建，正在处理...")


@app.get("/api/tasks/{task_id}", response_model=OCRResult)
async def get_task_status(task_id: str):
    """
    获取任务状态

    Args:
        task_id: 任务ID

    Returns:
        任务状态和结果
    """
    if task_id not in tasks:
        raise HTTPException(status_code=404, detail="任务不存在")

    task_info = tasks[task_id]

    # 如果任务完成，返回结果
    if task_id in results:
        return OCRResult(**results[task_id], task_id=task_id)

    # 否则返回当前状态
    return OCRResult(
        task_id=task_id, status=task_info["status"], progress=task_info["progress"]
    )


@app.get("/api/results/{task_id}")
async def get_results(task_id: str):
    """
    获取OCR结果

    Args:
        task_id: 任务ID

    Returns:
        OCR结果
    """
    if task_id not in results:
        raise HTTPException(status_code=404, detail="结果不存在")

    return results[task_id]


@app.delete("/api/tasks/{task_id}")
async def delete_task(task_id: str):
    """删除任务"""
    if task_id in tasks:
        del tasks[task_id]
    if task_id in results:
        del results[task_id]

    return {"message": "任务已删除"}


class TranslationRequest(BaseModel):
    text: str
    target_lang: str = "en"
    provider: str = "ollama"
    api_key: Optional[str] = None
    model: Optional[str] = None


class ConvertRequest(BaseModel):
    task_id: str
    target_format: str
    include_metadata: bool = True


@app.get("/api/stats")
async def get_stats():
    """获取系统统计"""
    return {
        "total_tasks": len(tasks),
        "completed_tasks": sum(1 for t in tasks.values() if t["status"] == "completed"),
        "queued_tasks": sum(1 for t in tasks.values() if t["status"] == "queued"),
        "processing_tasks": sum(
            1 for t in tasks.values() if t["status"] == "processing"
        ),
    }


@app.get("/api/files")
async def list_files():
    """列出所有上傳的檔案"""
    files = []
    if UPLOAD_DIR.exists():
        for f in UPLOAD_DIR.iterdir():
            if f.is_file():
                stat = f.stat()
                files.append(
                    {
                        "name": f.name,
                        "size": stat.st_size,
                        "created_at": stat.st_ctime,
                        "modified_at": stat.st_mtime,
                    }
                )
    return files


@app.delete("/api/files/{filename}")
async def delete_file(filename: str):
    """刪除檔案"""
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="檔案不存在")

    try:
        os.remove(file_path)
        return {"message": f"檔案 {filename} 已刪除"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"刪除失敗: {str(e)}")


@app.get("/api/files/{filename}/download")
async def download_file(filename: str):
    """下載檔案"""
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="檔案不存在")

    return FileResponse(
        path=file_path, filename=filename, media_type="application/octet-stream"
    )


@app.get("/api/plugins")
async def list_plugins():
    """列出所有插件"""
    return plugin_loader.list_plugins()


@app.post("/api/export")
async def export_results(
    data: dict,  # 接收 { "text": "...", "format": "docx", "filename": "..." }
):
    """
    將辨識結果匯出為指定格式的檔案
    """
    text = data.get("text", "")
    out_format = data.get("format", "docx").lower()
    base_name = data.get("filename", "ocr_result")

    # 確保輸出目錄存在
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    try:
        if out_format == "docx":
            from docx import Document

            doc = Document()
            doc.add_heading("OCR 辨識結果報告", 0)
            doc.add_paragraph(text)
            out_file = OUTPUT_DIR / f"{base_name}_{int(time.time())}.docx"
            doc.save(str(out_file))

        elif out_format == "xlsx":
            from openpyxl import Workbook

            wb = Workbook()
            ws = wb.active
            ws.title = "OCR Result"
            # 按行分割內容寫入 Excel
            lines = text.split("\n")
            for idx, line in enumerate(lines):
                ws.cell(row=idx + 1, column=1, value=line)
            out_file = OUTPUT_DIR / f"{base_name}_{int(time.time())}.xlsx"
            wb.save(str(out_file))

        elif out_format == "txt":
            # 新增純文字匯出
            out_file = OUTPUT_DIR / f"{base_name}_{int(time.time())}.txt"
            out_file.write_text(text, encoding="utf-8")

        else:
            return {"status": "error", "message": f"不支援的格式: {out_format}"}

        return {
            "status": "success",
            "download_url": f"/api/files/download/{out_file.name}?directory=output",
            "filename": out_file.name
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.get("/api/export-text/{task_id}")
async def export_text(task_id: str):
    """
    匯出任務的 OCR 結果為純文字檔案
    
    Args:
        task_id: 任務 ID
    
    Returns:
        FileResponse: 可下載的文字檔案
    """
    if task_id not in results:
        raise HTTPException(status_code=404, detail="任務不存在")
    
    task_result = results[task_id]
    if task_result.get("status") != "completed":
        raise HTTPException(status_code=400, detail="任務尚未完成")
    
    # 獲取文字內容
    text_content = task_result.get("results", {}).get("raw_result", "")
    
    # 生成文字檔案
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_file = OUTPUT_DIR / f"ocr_result_{task_id}_{int(time.time())}.txt"
    output_file.write_text(text_content, encoding="utf-8")
    
    # RFC 5987 編碼檔名
    final_filename = "ocr_result.txt"
    encoded_filename = quote(final_filename)

    return FileResponse(
        path=output_file,
        filename=final_filename,
        media_type="text/plain; charset=utf-8",
        headers={
            "Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"
        }
    )


@app.post("/api/translate")
async def translate_text(request: TranslationRequest):
    """
    使用 AI 翻譯文字
    
    Args:
        request: 翻譯請求物件
    
    Returns:
        翻譯後的文字
    """
    text = request.text
    target_lang = request.target_lang
    provider = request.provider
    api_key = request.api_key
    model = request.model
    
    print(f"[DEBUG] 開始處理翻譯請求. 提供商: {provider}, 目標語言: {target_lang}")
    print(f"[DEBUG] 文字長度: {len(text)} 字元")
    
    try:
        print("[DEBUG] 正在導入 LLM 客戶端...")
        from paddleocr_toolkit.llm.llm_client import create_llm_client
        print("[DEBUG] 導入成功")
        
        # 設定提示詞
        print("[DEBUG] 正在準備提示詞...")
        lang_map = {
            "en": "English",
            "zh-TW": "Traditional Chinese (繁體中文)",
            "zh-CN": "Simplified Chinese (简体中文)",
            "ja": "Japanese (日本語)",
            "ko": "Korean (한국어)",
            "es": "Spanish",
            "fr": "French",
            "de": "German"
        }
        
        target_language = lang_map.get(target_lang, target_lang)
        prompt = f"""請將以下文字翻譯成 {target_language}。
只需要輸出翻譯結果，不要有任何額外說明。

原文：
{text}

翻譯："""
        
        print(f"[DEBUG] 提示詞準備完成. 目標語言名稱: {target_language}")
        
        # 創建 LLM 客戶端
        print(f"[DEBUG] 正在創建 {provider} 客戶端...")
        kwargs = {}
        if api_key:
            kwargs["api_key"] = api_key
        if model:
            kwargs["model"] = model
        elif provider == "ollama":
            kwargs["model"] = "qwen2.5:7b"
            kwargs["base_url"] = "http://localhost:11434"
        
        llm = create_llm_client(provider=provider, **kwargs)
        print(f"[DEBUG] 客戶端創建成功")
        
        # 檢查服務是否可用
        print(f"[DEBUG] 正在檢查 {provider} 服務可用性...")
        if not llm.is_available():
            print(f"[ERROR] {provider} 服務檢查失敗 (不可用)")
            return {
                "status": "error",
                "message": f"{provider} 服務不可用，請確認已啟動相關服務"
            }
        
        print(f"[DEBUG] 服務可用性檢查通過")
        
        # 生成翻譯
        print(f"[DEBUG] 正在向 {provider} 發送生成請求 (此步驟可能需要一段時間)...")
        translated = llm.generate(prompt, temperature=0.3, max_tokens=4096)
        print(f"[DEBUG] 生成請求完成")
        
        if not translated:
            print(f"[WARNING] 翻譯結果為空")
            return {
                "status": "error",
                "message": "翻譯失敗，請重試"
            }
        
        print(f"[SUCCESS] 翻譯完成，長度: {len(translated)}")
        return {
            "status": "success",
            "translated_text": translated,
            "provider": provider,
            "target_lang": target_lang
        }
        
    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        print(f"[ERROR] 翻譯錯誤: {e}")
        print(error_detail)
        return {
            "status": "error",
            "message": f"翻譯失敗: {str(e)}"
        }


@app.post("/api/convert")
async def convert_format(request: ConvertRequest):
    """
    將 OCR 結果轉換為指定格式
    
    Args:
        request: 轉換請求物件
    
    Returns:
        FileResponse: 可下載的轉換檔案
    """
    task_id = request.task_id
    target_format = request.target_format
    include_metadata = request.include_metadata
    
    if task_id not in results:
        raise HTTPException(status_code=404, detail="任務不存在")
    
    task_result = results[task_id]
    if task_result.get("status") != "completed":
        raise HTTPException(status_code=400, detail="任務尚未完成")
    
    text_content = task_result.get("results", {}).get("raw_result", "")
    
    # 生成輸出檔案
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = int(time.time())
    
    try:
        from paddleocr_toolkit.utils.format_converter import FormatConverter
        
        converter = FormatConverter()
        
        if target_format == "docx":
            output_file = OUTPUT_DIR / f"ocr_result_{task_id}_{timestamp}.docx"
            converter.text_to_docx(text_content, str(output_file))
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        elif target_format == "xlsx":
            output_file = OUTPUT_DIR / f"ocr_result_{task_id}_{timestamp}.xlsx"
            converter.text_to_xlsx(text_content, str(output_file))
            media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        
        elif target_format == "pdf":
            output_file = OUTPUT_DIR / f"ocr_result_{task_id}_{timestamp}.pdf"
            converter.text_to_pdf_searchable(text_content, str(output_file))
            media_type = "application/pdf"
        
        elif target_format == "md":
            output_file = OUTPUT_DIR / f"ocr_result_{task_id}_{timestamp}.md"
            metadata = None
            if include_metadata:
                metadata = {
                    "date": time.strftime("%Y-%m-%d %H:%M:%S"),
                    "pages": task_result.get("results", {}).get("pages", 1),
                    "confidence": task_result.get("results", {}).get("confidence", 0.95)
                }
            converter.text_to_markdown(text_content, str(output_file), metadata)
            media_type = "text/markdown; charset=utf-8"
        
        else:  # txt (向後兼容)
            output_file = OUTPUT_DIR / f"ocr_result_{task_id}_{timestamp}.txt"
            output_file.write_text(text_content, encoding="utf-8")
            media_type = "text/plain; charset=utf-8"
        
        # RFC 5987 編碼檔名
        final_filename = f"ocr_result.{target_format}"
        encoded_filename = quote(final_filename)

        return FileResponse(
            path=output_file,
            filename=final_filename,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"
            }
        )
    
    except Exception as e:
        print(f"格式轉換失敗: {e}")
        raise HTTPException(status_code=500, detail=f"轉換失敗: {str(e)}")


@app.websocket("/ws/{task_id}")
async def websocket_endpoint(websocket: WebSocket, task_id: str):
    """
    WebSocket endpoint for real-time task updates
    """
    await manager.connect(websocket, task_id)
    try:
        while True:
            # Keep connection alive and listen for client messages if needed
            # For now, we just push updates from server
            data = await websocket.receive_text()
            # Optional: handle client commands
            if data == "ping":
                await manager.send_personal_message({"type": "pong"}, websocket)
    except WebSocketDisconnect:
        manager.disconnect(websocket, task_id)
    except Exception as e:
        print(f"WebSocket error: {e}")
        manager.disconnect(websocket, task_id)


if __name__ == "__main__":
    import uvicorn

    print(
        """
    ╔═══════════════════════════════════════════════════════╗
    ║                                                       ║
    ║     PaddleOCR Toolkit API Server                     ║
    ║     v3.3.0                                           ║
    ║                                                       ║
    ║     Dashboard: http://localhost:8000/                ║
    ║     Docs:      http://localhost:8000/docs            ║
    ║                                                       ║
    ╚═══════════════════════════════════════════════════════╝
    
    啟動伺服器...
    """
    )
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")


@app.get("/api/export-searchable-pdf/{task_id}")
async def export_searchable_pdf(task_id: str):
    """
    將 OCR 結果嵌入原始 PDF，生成可搜尋 PDF
    
    Args:
        task_id: 任務 ID
    
    Returns:
        FileResponse: 可下載的可搜尋 PDF 檔案
    """
    if task_id not in results:
        raise HTTPException(status_code=404, detail="任務不存在")
    
    task_result = results[task_id]
    if task_result.get("status") != "completed":
        raise HTTPException(status_code=400, detail="任務尚未完成")
    
    # 獲取原始檔案路徑
    original_file = task_result.get("file_path")
    if not original_file or not Path(original_file).exists():
        raise HTTPException(status_code=404, detail="原始檔案不存在")
    
    # 只支援 PDF 檔案
    if not original_file.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="僅支援 PDF 檔案")
    
    try:
        import fitz  # PyMuPDF
        
        # 獲取 OCR 結果
        ocr_results = task_result.get("results", {})
        raw_result = ocr_results.get("raw_result")
        
        # 開啟原始 PDF
        doc = fitz.open(original_file)
        
        # 增強版：使用 OCR 座標資訊進行精確定位
        if raw_result and isinstance(raw_result, list):
            # PaddleOCR 格式：[[page1_results], [page2_results], ...]
            for page_idx, page_ocr in enumerate(raw_result):
                if page_idx >= len(doc):
                    break
                
                page = doc[page_idx]
                page_rect = page.rect
                page_height = page_rect.height
                page_width = page_rect.width
                
                if not page_ocr:
                    continue
                
                # 處理每個文字區塊
                for line in page_ocr:
                    if not line or len(line) < 2:
                        continue
                    
                    # line[0] = 座標, line[1] = (文字, 信心度)
                    coords = line[0]  # [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
                    text_info = line[1]  # (text, confidence)
                    
                    if not coords or not text_info:
                        continue
                    
                    text = text_info[0] if isinstance(text_info, (tuple, list)) else str(text_info)
                    
                    if not text.strip():
                        continue
                    
                    # 計算邊界框
                    try:
                        x_coords = [p[0] for p in coords]
                        y_coords = [p[1] for p in coords]
                        x1, x2 = min(x_coords), max(x_coords)
                        y1, y2 = min(y_coords), max(y_coords)
                        
                        # 計算字體大小（基於高度）
                        height = y2 - y1
                        font_size = max(1, int(height * 0.8))  # 稍微小一點以確保不溢出
                        
                        # 插入透明文字（使用 insert_text 而非 insert_textbox）
                        # 位置使用左下角座標
                        try:
                            page.insert_text(
                                (x1, y2),
                                text,
                                fontsize=font_size,
                                color=(1, 1, 1),  # 白色（不可見）
                                overlay=False
                            )
                        except:
                            # 如果插入失敗，嘗試用更小的字體
                            page.insert_text(
                                (x1, y2),
                                text,
                                fontsize=max(1, font_size // 2),
                                color=(1, 1, 1),
                                overlay=False
                            )
                    except (IndexError, ValueError, TypeError) as e:
                        # 座標格式不正確，跳過這個區塊
                        continue
        else:
            # 降級方案：如果沒有座標資訊，使用舊方法
            raw_text = ocr_results.get("raw_result", "")
            if isinstance(raw_text, str) and raw_text and len(doc) > 0:
                page = doc[0]
                rect = page.rect
                try:
                    page.insert_textbox(
                        rect,
                        raw_text,
                        fontsize=1,
                        color=(1, 1, 1),
                        overlay=False
                    )
                except:
                    pass
        
        # 生成輸出檔案
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        output_file = OUTPUT_DIR / f"searchable_{task_id}_{int(time.time())}.pdf"
        doc.save(str(output_file))
        doc.close()
        
        # RFC 5987 編碼檔名
        final_filename = "searchable_ocr_result.pdf"
        encoded_filename = quote(final_filename)
        
        return FileResponse(
            path=output_file,
            filename=final_filename,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成可搜尋 PDF 失敗: {str(e)}")


# ==================== 批量處理 API ====================

@app.post("/api/batch-ocr")
async def batch_ocr(
    request: Request,
    files: list[UploadFile] = File(...),
    background_tasks: BackgroundTasks = None,
    mode: str = "basic",
    gemini_key: Optional[str] = None,
    claude_key: Optional[str] = None
):
    """
    批量處理多個 PDF/圖片檔案
    
    Args:
        files: 上傳的檔案列表
        mode: OCR 模式 (basic, hybrid, structure 等)
        gemini_key: Gemini API Key（可選）
        claude_key: Claude API Key（可選）
    
    Returns:
        {
            "batch_id": str,
            "task_ids": List[str],
            "total": int
        }
    """
    # 速率限制檢查（批量處理更嚴格）
    client_ip = request.client.host
    if not check_rate_limit(client_ip, RATE_LIMIT_BATCH):
        raise HTTPException(
            status_code=429,
            detail="批量處理請求過於頻繁，請稍後再試（限制：3 次/分鐘）"
        )
    
    if not files:
        raise HTTPException(status_code=400, detail="未選擇檔案")
    
    if len(files) > 20:
        raise HTTPException(status_code=400, detail="一次最多處理 20 個檔案")
    
    batch_id = str(uuid.uuid4())
    task_ids = []
    
    print(f"=" * 60)
    print(f"[BATCH] 批量處理請求")
    print(f"[BATCH] Batch ID: {batch_id}")
    print(f"[BATCH] 檔案數量: {len(files)}")
    print(f"[BATCH] OCR 模式: {mode}")
    print(f"=" * 60)
    
    try:
        for idx, file in enumerate(files):
            task_id = str(uuid.uuid4())
            
            # 儲存檔案
            file_extension = Path(file.filename).suffix.lower()
            file_path = UPLOAD_DIR / f"{task_id}{file_extension}"
            
            with open(file_path, "wb") as f:
                content = await file.read()
                f.write(content)
            
            print(f"[BATCH] 檔案 {idx+1}/{len(files)}: {file.filename} -> {file_path}")
            
            # 建立任務
            tasks[task_id] = {
                "status": "queued",
                "progress": 0,
                "created_at": datetime.now(),
                "batch_id": batch_id,
                "file_name": file.filename
            }
            
            # 添加背景任務
            background_tasks.add_task(
                process_ocr_task,
                task_id,
                str(file_path),
                mode,
                gemini_key,
                claude_key
            )
            
            task_ids.append(task_id)
        
        # 建立批量追蹤
        batches[batch_id] = {
            "task_ids": task_ids,
            "status": "processing",
            "completed": 0,
            "total": len(files),
            "created_at": datetime.now()
        }
        
        print(f"[BATCH] 批量任務已建立，開始處理")
        print(f"=" * 60)
        
        return {
            "batch_id": batch_id,
            "task_ids": task_ids,
            "total": len(files)
        }
        
    except Exception as e:
        print(f"[BATCH ERROR] {str(e)}")
        raise HTTPException(status_code=500, detail=f"批量處理失敗: {str(e)}")


@app.get("/api/batch/{batch_id}/status")
async def get_batch_status(batch_id: str):
    """
    查詢批量處理狀態
    
    Args:
        batch_id: 批量 ID
    
    Returns:
        {
            "batch_id": str,
            "total": int,
            "completed": int,
            "failed": int,
            "processing": int,
            "progress": float,
            "tasks": List[dict]
        }
    """
    if batch_id not in batches:
        raise HTTPException(status_code=404, detail="批量任務不存在")
    
    batch = batches[batch_id]
    task_ids = batch["task_ids"]
    
    # 收集所有任務狀態
    task_statuses = []
    completed = 0
    failed = 0
    processing = 0
    
    for task_id in task_ids:
        task = tasks.get(task_id, {})
        status = task.get("status", "unknown")
        
        if status == "completed":
            completed += 1
        elif status == "failed":
            failed += 1
        elif status in ["processing", "queued"]:
            processing += 1
        
        task_statuses.append({
            "task_id": task_id,
            "file_name": task.get("file_name", ""),
            "status": status,
            "progress": task.get("progress", 0)
        })
    
    # 計算整體進度
    progress = (completed / batch["total"] * 100) if batch["total"] > 0 else 0
    
    # 更新批量狀態
    if completed + failed == batch["total"]:
        batch["status"] = "completed"
    
    return {
        "batch_id": batch_id,
        "total": batch["total"],
        "completed": completed,
        "failed": failed,
        "processing": processing,
        "progress": round(progress, 2),
        "tasks": task_statuses
    }


@app.get("/api/batch/{batch_id}/results")
async def get_batch_results(batch_id: str):
    """
    獲取批量處理結果
    
    Args:
        batch_id: 批量 ID
    
    Returns:
        {
            "batch_id": str,
            "results": List[dict]
        }
    """
    if batch_id not in batches:
        raise HTTPException(status_code=404, detail="批量任務不存在")
    
    batch = batches[batch_id]
    task_ids = batch["task_ids"]
    
    # 收集所有結果
    batch_results = []
    for task_id in task_ids:
        if task_id in results:
            result = results[task_id]
            batch_results.append({
                "task_id": task_id,
                "file_name": tasks.get(task_id, {}).get("file_name", ""),
                "status": result.get("status", "unknown"),
                "result": result.get("results", {})
            })
    
    return {
        "batch_id": batch_id,
        "results": batch_results
    }


@app.delete("/api/batch/{batch_id}")
async def delete_batch(batch_id: str):
    """
    刪除批量任務及相關資料
    
    Args:
        batch_id: 批量 ID
    """
    if batch_id not in batches:
        raise HTTPException(status_code=404, detail="批量任務不存在")
    
    batch = batches[batch_id]
    task_ids = batch["task_ids"]
    
    # 刪除所有相關任務和結果
    for task_id in task_ids:
        if task_id in tasks:
            del tasks[task_id]
        if task_id in results:
            del results[task_id]
    
    # 刪除批量記錄
    del batches[batch_id]
    
    return {"message": f"批量任務 {batch_id} 已刪除"}
