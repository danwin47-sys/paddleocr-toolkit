# -*- coding: utf-8 -*-
"""
PaddleOCR Toolkit - 格式轉換處理器

支援將 OCR 結果轉換為多種檔案格式
"""

from pathlib import Path
from typing import List, Optional, Dict, Any
import time


class FormatConverter:
    """檔案格式轉換處理器"""
    
    @staticmethod
    def text_to_docx(text: str, output_path: str) -> bool:
        """
        純文字轉 DOCX
        
        Args:
            text: 文字內容
            output_path: 輸出檔案路徑
        
        Returns:
            bool: 是否成功
        """
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading('OCR 辨識結果', 0)
            
            # 按段落分割並添加
            paragraphs = text.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"轉換 DOCX 失敗: {e}")
            return False
    
    @staticmethod
    def text_to_xlsx(text: str, output_path: str) -> bool:
        """
        純文字轉 Excel（每行一個儲存格）
        
        Args:
            text: 文字內容
            output_path: 輸出檔案路徑
        
        Returns:
            bool: 是否成功
        """
        try:
            from openpyxl import Workbook
            
            wb = Workbook()
            ws = wb.active
            ws.title = "OCR Result"
            
            # 添加標題
            ws['A1'] = 'OCR 辨識結果'
            ws['A1'].font = ws['A1'].font.copy(bold=True, size=14)
            
            # 按行分割並寫入
            lines = text.split('\n')
            for idx, line in enumerate(lines, 2):  # 從第2行開始
                ws.cell(row=idx, column=1, value=line)
            
            # 自動調整列寬
            ws.column_dimensions['A'].width = 80
            
            wb.save(output_path)
            return True
        except Exception as e:
            print(f"轉換 XLSX 失敗: {e}")
            return False
    
    @staticmethod
    def text_to_markdown(
        text: str, 
        output_path: str, 
        metadata: Optional[Dict[str, Any]] = None
    ) -> bool:
        """
        轉換為 Markdown 格式
        
        Args:
            text: 文字內容
            output_path: 輸出檔案路徑
            metadata: 元數據（日期、頁數、信心度等）
        
        Returns:
            bool: 是否成功
        """
        try:
            md_content = "# OCR 辨識結果\n\n"
            
            # 添加元數據
            if metadata:
                md_content += "## 📊 辨識資訊\n\n"
                if 'date' in metadata:
                    md_content += f"- **辨識日期**: {metadata['date']}\n"
                if 'pages' in metadata:
                    md_content += f"- **頁數**: {metadata['pages']}\n"
                if 'confidence' in metadata:
                    confidence_pct = int(metadata['confidence'] * 100)
                    md_content += f"- **信心度**: {confidence_pct}%\n"
                md_content += "\n---\n\n"
            
            md_content += "## 📄 辨識內容\n\n"
            md_content += text
            
            Path(output_path).write_text(md_content, encoding='utf-8')
            return True
        except Exception as e:
            print(f"轉換 Markdown 失敗: {e}")
            return False
    
    @staticmethod
    def text_to_pdf_searchable(
        text: str,
        output_path: str,
        image_path: Optional[str] = None,
        ocr_results: Optional[List] = None
    ) -> bool:
        """
        使用現有 PDFGenerator 生成可搜尋 PDF
        
        Args:
            text: 文字內容
            output_path: 輸出檔案路徑
            image_path: 原始圖片路徑（可選）
            ocr_results: OCR 結果列表（可選，用於精確定位）
        
        Returns:
            bool: 是否成功
        """
        try:
            from paddleocr_toolkit.core.pdf_generator import PDFGenerator
            
            # 如果有原始圖片和 OCR 結果，使用完整方法
            if image_path and ocr_results:
                generator = PDFGenerator(output_path)
                generator.add_page(image_path, ocr_results)
                return generator.save()
            
            # 否則創建純文字 PDF（簡化版）
            import fitz  # PyMuPDF
            
            doc = fitz.open()
            page = doc.new_page(width=595, height=842)  # A4 尺寸
            
            # 添加文字
            rect = fitz.Rect(50, 50, 545, 792)
            page.insert_textbox(
                rect,
                text,
                fontsize=12,
                fontname="helv",
                color=(0, 0, 0)
            )
            
            doc.save(output_path)
            doc.close()
            return True
            
        except Exception as e:
            print(f"轉換 PDF 失敗: {e}")
            return False
